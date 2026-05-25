"""
Steps AI Resume Text Extraction Service.

This module provides text decoding capabilities for layout-aware document formats
including PDF (using PyMuPDF) and Microsoft Word DOCX (using python-docx).
"""

import fitz  # PyMuPDF
import docx
import io
from fastapi import HTTPException
import logging

logger = logging.getLogger(__name__)

def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """
    Extracts raw text content from PDF binary streams using PyMuPDF (fitz).
    """
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text_content = []
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text_content.append(page.get_text())
        doc.close()
        return "\n".join(text_content).strip()
    except Exception as e:
        logger.error(f"Failed to parse PDF: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=400,
            detail=f"Failed to extract text from PDF document: {str(e)}"
        )

def extract_text_from_docx(docx_bytes: bytes) -> str:
    """
    Extracts raw text content from Microsoft Word DOCX binary streams using python-docx.
    """
    try:
        doc = docx.Document(io.BytesIO(docx_bytes))
        text_content = []
        for paragraph in doc.paragraphs:
            text_content.append(paragraph.text)
        # Also extract text from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_content.append(cell.text)
        return "\n".join(text_content).strip()
    except Exception as e:
        logger.error(f"Failed to parse DOCX: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=400,
            detail=f"Failed to extract text from Microsoft Word DOCX document: {str(e)}"
        )

def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Decodes raw document text based on file extensions (.pdf or .docx).
    """
    lower_filename = filename.lower()
    if lower_filename.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
    elif lower_filename.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
    else:
        raise HTTPException(
            status_code=400,
            detail="Unsupported document format. Please upload a valid PDF (.pdf) or Microsoft Word (.docx) document."
        )

    if not text:
        raise HTTPException(
            status_code=400,
            detail="The uploaded document contains no readable text content."
        )

    return text
